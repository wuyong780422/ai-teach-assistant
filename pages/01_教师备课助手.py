import streamlit as st
from docx import Document
from openai import OpenAI
import os
import re
from io import BytesIO

# 页面基础配置
st.set_page_config(
    page_title="AI备课助手",
    page_icon="📖",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# ========== 以下全部是你原来的备课代码，完整保留 ==========
class AILessonPlanner:
    def __init__(self):
        self.template_path = self._get_template_path("lesson_template.docx")
        self.ai_client = self._init_ai_client()
        self.cell_pos = {
            "subject": [(4, 1)],
            "grade": [(4, 6)],
            "teach_content": [(6, 2)],
            "total_class": [(6, 10)],
            "knowledge_base": [(8, 8), (8, 9), (8, 10), (8, 11)],
            "cognitive_ability": [(9, 8), (9, 9), (9, 10), (9, 11)],
            "learning_features": [(10, 8), (10, 9), (10, 10), (10, 11)],
            "knowledge_goal": [(11, 8), (11, 9), (11, 10), (11, 11)],
            "ability_goal": [(12, 8), (12, 9), (12, 10), (12, 11)],
            "quality_goal": [(13, 8), (13, 9), (13, 10), (13, 11)],
            "key_points": [(14, 2), (14, 3)],
            "difficult_points": [(15, 2), (15, 3)],
            "teaching_methods": [(16, 2), (16, 3)],
            "teaching_prepare": [(17, 2), (17, 3)],
            "pre_content": [(21, 2), (21, 3), (21, 4)],
            "pre_teacher": [(21, 5), (21, 6), (21, 7)],
            "pre_student": [(21, 10), (21, 11), (21, 12)],
            "class_content": [(25, 2), (25, 3), (25, 4)],
            "class_teacher": [(25, 5), (25, 6), (25, 7)],
            "class_student": [(25, 10), (25, 11), (25, 12)],
            "after_content": [(29, 2), (29, 3), (29, 4)],
            "after_teacher": [(29, 5), (29, 6), (29, 7)],
            "after_student": [(29, 10), (29, 11), (29, 12)],
            "blackboard": [(31, 9), (31, 10), (31, 11), (31, 12)]
        }
    def _get_template_path(self, filename):
        base_path = os.path.dirname(os.path.abspath(__file__))
        # 模板在【根目录】，回退一级目录
        base_path = os.path.dirname(base_path)
        template_path = os.path.join(base_path, filename)
        if not os.path.exists(template_path):
            st.error(f"未找到模板文件：{template_path}，请确认lesson_template.docx在程序根目录")
            st.stop()
        return template_path
    def _init_ai_client(self):
        try:
            api_key = st.secrets["deepseek_api_key"]
            base_url = st.secrets.get("deepseek_base_url", "https://api.deepseek.com/v1")
            return OpenAI(
                api_key=api_key,
                base_url=base_url
            )
        except Exception as e:
            st.error(f"AI接口初始化失败，请检查后台密钥配置：{str(e)}")
            st.stop()
    def _build_prompt(self, user_input):
        return f"""你是中职专业课教师，撰写{user_input['grade']}{user_input['major']}《{user_input['topic']}》教案，若为数学专业，主题内容为平面几何相关，则贴合苗族服饰案例来设计教案。硬性规则：
        1、**所有正文内容严禁出现任何字段文字（如xx值：、数字+字段名全部禁止写入正文），字段名只在每行开头用于分类，正文不能携带字段标识前缀**；
        2、课前、课后内容自由撰写无分段；课中教学内容、课中教师活动、课中学生活动都必须包含：
        1）引入4min 2）探究8min 3）讲解10min 4）拓展10min 5）评价8min
        所有内容合并为单独一行（禁止换行）；【强制字数+内容硬性要求，必须严格执行，严禁简写】
        课中教学内容：每个环节用20字左右精简概括核心知识点，控制在15-25字；
        课中教师活动：
        1）2）3）三个环节：写清具体教学动作，字数严格控制在35-45字之间，内容完整详实，严禁只写短语；
        4）拓展10min环节：先写原教学内容（35-45字），再添加加一段约40字的课程思政内容，格式为“课程思政：XXX”，拓展环节总字数控制在40-50字之间，思政点需结合本课专业内容，如职业规范、网络安全伦理、工匠精神、家国情怀等，自然融入不生硬；
        5）评价8min：先写原教学内容（35-45字），再添加一段约40字的课堂小结内容，格式为“课堂小结：XXX”，梳理本节课核心知识点，字数控制在35-45字；
        5）评价8min：添加加一段约40字的课堂小结内容之后，再添加一段约40字的作业布置内容，格式为“作业布置：XXX”，布置课后实操类作业，字数控制在35-45字；
        课中学生活动：每个环节必须写清具体互动行为（如举手回答、分组讨论、上机操作、判断分析等），字数严格控制在25-35字之间，贴合课堂实际；
        三项内容彼此独立，严禁互相摘抄、复用文本，每个环节字数偏差不得超过20个字，禁止过度简写或长篇堆砌。
        3、课中教学内容、课中教师活动、课中学生活动三项内容互相完全独立，禁止交叉摘抄；
        4、课后学生活动必须独立设计家庭实操作业，严禁照搬课中任何文字；
        5、全字段通用：课中、课后内容禁止摘抄学情、三维目标、重难点、教学准备、课前原文；
        6、【重中之重：20.板书设计值**必须单独独占一行开始，绝对不能拼接在上一字段末尾，独立创作；严禁出现引入、探究、精解、拓展、评价、min、分组、演示、1）、2）、3）、4）、5）等课中步骤词汇，禁止照搬任何课中文字，固定3～5条短提纲，每条≤10个汉字，只提炼课程核心知识点短句，不许书写课堂步骤**】；
        7、全文无*、无多余特殊符号，禁用万能空洞套话；
        8、【专项】学情3项+三维目标共6处，**禁用理解、掌握、了解、学会等模糊词语，全部使用能说出、能列举、能实操等可量化、可检测的可测评句式；六项内容各自独立，侧重点互不相同，严禁整段互相复制复用**；
        输出严格20行：字段名：正文
        1.知识与技能基础值：
        2.认知和实践能力值：
        3.学生学习特点值：
        4.知识目标值：
        5.能力目标值：
        6.素养目标值：
        7.教学重点值：
        8.教学难点值：
        9.教学方法值：
        10.教学准备值：
        11.课前教学内容值：
        12.课前教师活动值：
        13.课前学生活动值：
        14.课中教学内容值：
        15.课中教师活动值：
        16.课中学生活动值：
        17.课后教学内容值：
        18.课后教师活动值：
        19.课后学生活动值：
        20.板书设计值："""
    def _get_preset_content(self, field_name, user_input):
        return "AI接口异常，无法获取教案内容"
    def _extract_ai_value(self, ai_content, field_name, user_input):
        black_words = [
            "知识与技能基础", "认知和实践能力", "学生学习特点",
            "知识目标", "能力目标", "素养目标", "教学重点",
            "教学难点", "教学方法", "教学准备",
            "课前教学内容", "课前教师活动", "课前学生活动"
        ]
        ai_text = str(ai_content).strip()
        if len(ai_text) < 10:
            return ""
        lines = [l.strip() for l in ai_text.splitlines() if l.strip()]
        target_key = [field_name, field_name.replace("值", ""), field_name[:6]]
        res = ""
        start_idx = -1
        for i, line in enumerate(lines):
            for kw in target_key:
                if kw in line and ("：" in line or ":" in line):
                    start_idx = i
                    break
            if start_idx != -1:
                break
        if start_idx == -1:
            return ""
        if field_name == "板书设计值":
            first_line = lines[start_idx]
            split_char = "：" if "：" in first_line else ":"
            res = first_line.split(split_char, 1)[1].strip()
            for j in range(start_idx + 1, len(lines)):
                next_line = lines[j]
                if re.match(r"^\d+\.", next_line):
                    break
                res += "、" + next_line.strip()
        else:
            line = lines[start_idx]
            split_char = "：" if "：" in line else ":"
            temp_text = line.split(split_char, 1)[1].strip()
            if any(bw in temp_text for bw in black_words):
                return ""
            res = temp_text
        res = re.sub(r"\d+\..*?[：:]", "", res)
        res = res.replace("*", "").strip()
        if field_name == "板书设计值":
            res = res.lstrip("、")
            if len(res) > 120:
                res = res[:120]
        major_str = str(user_input.get("major", ""))
        res = res.replace("{major}", major_str)
        return res
    def _generate_ai_content(self, user_input):
        ai_content = ""
        if self.ai_client is not None:
            try:
                msg_list = [{"role": "user", "content": self._build_prompt(user_input)}]
                resp = self.ai_client.chat.completions.create(
                    model="deepseek-chat",
                    messages=msg_list,
                    temperature=0.32,
                    max_tokens=7000,
                    timeout=25
                )
                ai_content = resp.choices[0].message.content.strip()
            except Exception as err:
                print("AI调用异常：", err)
        return ai_content
    def _safe_set_cell(self, table, pos_list, text):
        for (row, col) in pos_list:
            try:
                if 0 <= row < len(table.rows) and 0 <= col < len(table.columns):
                    table.cell(row, col).text = text
                    return True
            except Exception as e:
                print(f"单元格({row},{col})填充异常：{e}")
        return False
    def _fill_template(self, ai_content, user_input):
        doc = Document(self.template_path)
        if len(doc.tables) == 0:
            st.error("模板表格不存在，请检查模板文件")
            return None
        table = doc.tables[0]
        self._safe_set_cell(table, self.cell_pos["subject"], user_input["subject"])
        self._safe_set_cell(table, self.cell_pos["grade"], user_input["grade"])
        self._safe_set_cell(table, self.cell_pos["teach_content"], user_input["topic"])
        self._safe_set_cell(table, self.cell_pos["total_class"], "2课时")
        kb = self._extract_ai_value(ai_content, "知识与技能基础值", user_input)
        ca = self._extract_ai_value(ai_content, "认知和实践能力值", user_input)
        lf = self._extract_ai_value(ai_content, "学生学习特点值", user_input)
        for pos in self.cell_pos["knowledge_base"]:
            self._safe_set_cell(table, [pos], kb)
        for pos in self.cell_pos["cognitive_ability"]:
            self._safe_set_cell(table, [pos], ca)
        for pos in self.cell_pos["learning_features"]:
            self._safe_set_cell(table, [pos], lf)
        kg = self._extract_ai_value(ai_content, "知识目标值", user_input)
        ag = self._extract_ai_value(ai_content, "能力目标值", user_input)
        qg = self._extract_ai_value(ai_content, "素养目标值", user_input)
        for pos in self.cell_pos["knowledge_goal"]:
            self._safe_set_cell(table, [pos], kg)
        for pos in self.cell_pos["ability_goal"]:
            self._safe_set_cell(table, [pos], ag)
        for pos in self.cell_pos["quality_goal"]:
            self._safe_set_cell(table, [pos], qg)
        kp = self._extract_ai_value(ai_content, "教学重点值", user_input)
        dp = self._extract_ai_value(ai_content, "教学难点值", user_input)
        tm = self._extract_ai_value(ai_content, "教学方法值", user_input)
        tp = self._extract_ai_value(ai_content, "教学准备值", user_input)
        self._safe_set_cell(table, self.cell_pos["key_points"], kp)
        self._safe_set_cell(table, self.cell_pos["difficult_points"], dp)
        self._safe_set_cell(table, self.cell_pos["teaching_methods"], tm)
        self._safe_set_cell(table, self.cell_pos["teaching_prepare"], tp)
        pre_c = self._extract_ai_value(ai_content, "课前教学内容值", user_input)
        pre_t = self._extract_ai_value(ai_content, "课前教师活动值", user_input)
        pre_s = self._extract_ai_value(ai_content, "课前学生活动值", user_input)
        self._safe_set_cell(table, self.cell_pos["pre_content"], pre_c)
        for pos in self.cell_pos["pre_teacher"]:
            self._safe_set_cell(table, [pos], pre_t)
        for pos in self.cell_pos["pre_student"]:
            self._safe_set_cell(table, [pos], pre_s)
        cla_c = self._extract_ai_value(ai_content, "课中教学内容值", user_input)
        cla_t = self._extract_ai_value(ai_content, "课中教师活动值", user_input)
        cla_s = self._extract_ai_value(ai_content, "课中学生活动值", user_input)
        self._safe_set_cell(table, self.cell_pos["class_content"], cla_c)
        for pos in self.cell_pos["class_teacher"]:
            self._safe_set_cell(table, [pos], cla_t)
        for pos in self.cell_pos["class_student"]:
            self._safe_set_cell(table, [pos], cla_s)
        aft_c = self._extract_ai_value(ai_content, "课后教学内容值", user_input)
        aft_t = self._extract_ai_value(ai_content, "课后教师活动值", user_input)
        aft_s = self._extract_ai_value(ai_content, "课后学生活动值", user_input)
        self._safe_set_cell(table, self.cell_pos["after_content"], aft_c)
        for pos in self.cell_pos["after_teacher"]:
            self._safe_set_cell(table, [pos], aft_t)
        for pos in self.cell_pos["after_student"]:
            self._safe_set_cell(table, [pos], aft_s)
        bb = self._extract_ai_value(ai_content, "板书设计值", user_input)
        for pos in self.cell_pos["blackboard"]:
            self._safe_set_cell(table, [pos], bb)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    def generate_lesson_plan(self, user_input):
        if not user_input["topic"]:
            st.warning("请填写授课主题！")
            return None
        with st.spinner("AI正在独立原创教案，请稍候..."):
            ai_txt = self._generate_ai_content(user_input)
            if not ai_txt:
                st.error("AI生成失败，请检查网络或密钥配置")
                return None
            file_buf = self._fill_template(ai_txt, user_input)
            if file_buf is None:
                return None
            file_name = f"{user_input['grade']}_{user_input['major']}_{user_input['topic']}_教案.docx"
            return file_buf, file_name

# 页面UI
st.title("📖 备课助手")
st.caption("中职专业课教案一键生成 · 手机电脑通用")

with st.form("lesson_form"):
    col1, col2 = st.columns(2)
    with col1:
        major = st.text_input("专业", value="计算机应用")
        grade = st.text_input("年级", value="高一")
    with col2:
        subject = st.text_input("学科", value="信息技术")
        topic = st.text_input("授课主题", value="计算机病毒")
    submit_btn = st.form_submit_button("生成完整教案", use_container_width=True)

if submit_btn:
    user_input = {
        "major": major.strip() or "计算机应用",
        "subject": subject.strip() or "信息技术",
        "grade": grade.strip() or "高一",
        "topic": topic.strip() or "计算机病毒"
    }
    planner = AILessonPlanner()
    result = planner.generate_lesson_plan(user_input)
    if result:
        file_buf, file_name = result
        st.success("✅ 教案生成成功！")
        st.download_button(
            label="📥 下载Word教案",
            data=file_buf,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# 返回教师首页
st.divider()
st.page_link("pages/00_教师入口首页.py", label="← 返回教师入口功能", use_container_width=True)