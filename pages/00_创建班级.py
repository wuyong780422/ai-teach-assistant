import streamlit as st
import pandas as pd
import os
from data_tools import add_class, get_all_class, get_students_by_class, get_class_all_info, get_question_list, \
    STUDENT_CSV, get_all_papers
from data_tools import init_point_csv, get_class_point, add_class_point, add_group_all_point, get_class_group_map, \
    set_student_group, get_group_point_stat
from data_tools import get_class_resources, RESOURCE_FILE_DIR
from io import BytesIO
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="班级管理", layout="centered", initial_sidebar_state="collapsed")

# 初始化会话变量
if "random_pick_stu" not in st.session_state:
    st.session_state["random_pick_stu"] = ""
if "pick_group" not in st.session_state:
    st.session_state["pick_group"] = ""

st.markdown("<h1 style='text-align:center;'>课堂教学管理</h1>", unsafe_allow_html=True)
st.divider()

SCORE_CSV = "score_data.csv"
ENCODING = "utf-8-sig"
init_point_csv()

# 可在线预览的格式白名单
IMG_EXT = {"jpg", "jpeg", "png", "gif"}
VIDEO_EXT = {"mp4", "webm"}

# 1. 新建班级
st.subheader("一、班级管理")
input_new_class = st.text_input("输入班级名称（如：24级计算机1班）")
if st.button("确认创建班级", type="secondary"):
    if input_new_class.strip() == "":
        st.warning("班级名称不可为空！")
    else:
        flag, msg = add_class(input_new_class.strip())
        if flag:
            st.success(f"创建成功！验证码：{msg}")
        else:
            st.error(msg)

st.divider()

# 2. 学生分组、抽选加分模块
st.subheader("二、课堂加分操作区")
all_class = get_all_class()
class_dict = get_class_all_info()
selected_class = ""

if len(all_class) == 0:
    st.info("暂无班级，请先创建班级")
else:
    selected_class = st.selectbox("选择操作班级", all_class)
    st.info(f"当前班级验证码：{class_dict[selected_class]}")
    stu_list = get_students_by_class(selected_class)
    group_map = get_class_group_map(selected_class)
    group_list = list(group_map.keys())

    st.write(f"本班学生列表：")
    if len(stu_list) == 0:
        st.info("暂无学生，请学生登录加入班级")
    else:
        row_count = (len(stu_list) + 7) // 8
        idx = 0
        for _ in range(row_count):
            cols = st.columns(8)
            for col in cols:
                if idx < len(stu_list):
                    name = stu_list[idx]
                    if name == st.session_state["random_pick_stu"]:
                        col.button(name, use_container_width=True, type="primary")
                    else:
                        col.button(name, use_container_width=True)
                    idx += 1

        st.divider()
        # ① 随机抽取1名学生
        if st.button("随机抽取组和学生", type="primary", use_container_width=True):
            import random
            pick_name = random.choice(stu_list)
            st.session_state["random_pick_stu"] = pick_name
            df_stu = pd.read_csv(STUDENT_CSV, encoding="utf-8-sig")
            df_stu["group"] = df_stu["group"].fillna("").astype(str)
            target_row = df_stu[(df_stu["class_name"] == selected_class) & (df_stu["student_name"] == pick_name)]
            pick_group = ""
            if len(target_row) > 0:
                group_val = target_row.iloc[0]["group"]
                if group_val not in ["", "nan"]:
                    if group_val.endswith(".0"):
                        group_val = group_val[:-2]
                    pick_group = group_val
            st.session_state["pick_group"] = pick_group
            st.rerun()

        # ② 高亮显示被抽组和学生
        if st.session_state["random_pick_stu"] != "":
            picked_name = st.session_state["random_pick_stu"]
            df_stu = pd.read_csv(STUDENT_CSV, encoding="utf-8-sig")
            df_stu["group"] = df_stu["group"].fillna("").astype(str)
            target_row = df_stu[(df_stu["class_name"] == selected_class) & (df_stu["student_name"] == picked_name)]
            group_text = ""
            pick_group = ""
            if len(target_row) > 0:
                group_val = target_row.iloc[0]["group"]
                if group_val not in ["", "nan"]:
                    if group_val.endswith(".0"):
                        group_val = group_val[:-2]
                    group_text = f"第{group_val}组--"
                    pick_group = group_val
            show_text = f"本次抽中组和学生：{group_text}{picked_name}"
            st.markdown(f"""
            <div style="background:#165DFF;color:#fff;padding:20px;border-radius:12px;text-align:center;margin:15px 0;">
            <h2>{show_text}</h2>
            </div>
            """, unsafe_allow_html=True)

        # ③单人加分、④小组下拉、⑤整组加分、⑥清空 四栏并排
        col_c1, col_c2, col_c3, col_c4 = st.columns([2, 2, 2, 2])
        with col_c1:
            picked = st.session_state["random_pick_stu"]
            if picked != "":
                if st.button(f"为【{picked}】+1课堂分", use_container_width=True):
                    flag, tip = add_class_point(selected_class, picked)
                    st.success(tip)
                    st.rerun()
            else:
                st.button("为选中学生+1课堂分", disabled=True, use_container_width=True)
        with col_c2:
            default_group = st.session_state.get("pick_group", "")
            select_index = 0
            if default_group in group_list:
                select_index = group_list.index(default_group)
            sel_group = st.selectbox("选择小组批量加分", group_list, index=select_index)
        with col_c3:
            if len(group_list) > 0:
                if st.button("整组全员+1分", use_container_width=True):
                    flag, tip = add_group_all_point(selected_class, sel_group)
                    st.success(tip)
                    st.rerun()
            else:
                st.button("整组全员+1分", disabled=True, use_container_width=True)
        with col_c4:
            if st.button("清空抽取结果", use_container_width=True):
                st.session_state["random_pick_stu"] = ""
                st.session_state["pick_group"] = ""
                st.rerun()

        st.divider()
        # ⑦ 给学生分配小组
        st.subheader("给学生分配小组")
        c1, c2 = st.columns(2)
        sel_student = c1.selectbox("选择学生", stu_list)
        group_input = c2.text_input("小组编号（1/2/3...）", value="1")
        if st.button("保存小组设置"):
            flag, tip = set_student_group(selected_class, sel_student, group_input.strip())
            if flag:
                st.success(tip)
                st.rerun()
            else:
                st.error(tip)

        st.divider()
        # 个人积分排行榜
        st.subheader("三、学生个人积分排行榜（总分降序）")
        point_df = get_class_point(selected_class)
        if point_df.empty:
            st.info("暂无积分记录，加分后自动生成")
        else:
            point_df = point_df.sort_values("total_point", ascending=False, ignore_index=True)
            st.dataframe(point_df, use_container_width=True)

        st.divider()
        # 小组积分柱状图
        st.subheader("四、小组积分统计图表")
        group_stat = get_group_point_stat(selected_class)
        if group_stat.empty:
            st.info("请先给学生分配小组后查看统计")
        else:
            chart1, chart2 = st.columns(2)
            with chart1:
                st.subheader("各组总积分")
                st.bar_chart(group_stat, x="小组", y="小组总积分", height=320, use_container_width=True)
            with chart2:
                st.subheader("各组人均积分")
                st.bar_chart(group_stat, x="小组", y="人均积分", height=320, use_container_width=True)
            st.dataframe(group_stat, use_container_width=True)

# ========== 导出函数已提前移动到此处，保证调用前已定义 ==========
# ========== 导出1：学生成绩明细 Excel（与页面表格完全一致） ==========
def export_score_detail_excel(df, class_name, paper_name):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="学生成绩明细", index=False)
        ws = writer.sheets["学生成绩明细"]
        # 表头加粗
        for cell in ws[1]:
            cell.font = cell.font.copy(bold=True)
        # 自适应列宽
        for col in ws.columns:
            col_letter = col[0].column_letter
            max_len = max(len(str(c.value)) for c in col if c.value)
            ws.column_dimensions[col_letter].width = max_len * 1.2 + 2
    output.seek(0)
    return output

# ========== 导出2：试卷分析报告 Word（与页面分析数据完全一致） ==========
def export_paper_analysis_word(stat_data, segment_data, question_df, class_name, paper_name):
    doc = Document()
    # 标题
    title = doc.add_heading(f"{class_name} - {paper_name} 试卷分析报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    # 一、整体情况总览
    doc.add_heading("一、考试整体情况", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["参考人数", "平均分", "最高分", "最低分", "及格率", "优秀率"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    row_cells = table.add_row().cells
    row_cells[0].text = str(stat_data["total_count"])
    row_cells[1].text = f"{stat_data['avg_score']:.2f}"
    row_cells[2].text = f"{stat_data['max_score']:.0f}"
    row_cells[3].text = f"{stat_data['min_score']:.0f}"
    row_cells[4].text = f"{stat_data['pass_rate']:.1f}%"
    row_cells[5].text = f"{stat_data['excellent_rate']:.1f}%"
    doc.add_paragraph()
    # 二、分数段分布（内嵌柱状图）
    doc.add_heading("二、分数段人数分布", level=2)
    # 解决中文乱码
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(6, 3.5))
    segments = list(segment_data.keys())
    counts = list(segment_data.values())
    bars = ax.bar(segments, counts, color="#165DFF")
    ax.set_title("各分数段人数分布", fontsize=12)
    ax.set_ylabel("人数", fontsize=10)
    ax.set_xlabel("分数段", fontsize=10)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f"{int(height)}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    img_buf = BytesIO()
    plt.savefig(img_buf, format="png", dpi=150)
    img_buf.seek(0)
    plt.close()
    doc.add_picture(img_buf, width=Inches(5.5))
    doc.add_paragraph()
    # 三、逐题质量分析
    doc.add_heading("三、逐题质量分析", level=2)
    q_table = doc.add_table(rows=1, cols=4)
    q_table.style = "Table Grid"
    q_hdr = q_table.rows[0].cells
    q_headers = ["题号", "题型", "分值", "全班正确率"]
    for i, h in enumerate(q_headers):
        q_hdr[i].text = h
        q_hdr[i].paragraphs[0].runs[0].bold = True
    for _, row in question_df.iterrows():
        row_cells = q_table.add_row().cells
        row_cells[0].text = str(row["题号"])
        row_cells[1].text = str(row["题型"])
        row_cells[2].text = str(row["分值"])
        row_cells[3].text = f"{row['正确率']:.1f}%"
        # 正确率低于60%自动标红，和页面高亮逻辑一致
        if row["正确率"] < 60:
            run = row_cells[3].paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
    doc.add_paragraph()
    doc.add_heading("四、教学反思与改进措施", level=2)
    doc.add_paragraph("________________________________________________________________")
    doc.add_paragraph("________________________________________________________________")
    doc.add_paragraph("________________________________________________________________")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ===================== 五、班级试卷成绩分析 =====================
st.divider()
st.subheader("五、班级试卷成绩分析")
paper_list = get_all_papers()
if len(paper_list) == 0:
    st.info("暂无可用试卷，请先在试卷管理中创建试卷")
else:
    paper_map = {p["paper_id"]: p["paper_name"] for p in paper_list}
    select_test = st.selectbox(
        "选择试卷",
        list(paper_map.keys()),
        format_func=lambda x: paper_map[x]
    )
    q_list = get_question_list(select_test)
    if not selected_class:
        st.info("请先选择上方班级")
    elif len(q_list) == 0:
        st.warning("题库question_bank.csv暂无该试卷的题目")
    else:
        st.write(f"班级：{selected_class} | 试卷：{paper_map[select_test]}")
        btn1, btn2 = st.columns(2)
        with btn1:
            show_stat = st.button("加载班级统计分析", type="primary")
        with btn2:
            show_detail = st.button("加载全部成绩明细")

        df_score = pd.DataFrame()
        if os.path.exists(SCORE_CSV):
            df_score = pd.read_csv(SCORE_CSV, encoding=ENCODING)
        class_score = pd.DataFrame()
        if not df_score.empty and "test_id" in df_score.columns:
            class_score = df_score[
                (df_score["class_name"] == selected_class) & (df_score["test_id"] == select_test)].copy()

        # 成绩明细 + 新增Excel导出按钮
        if show_detail:
            st.divider()
            st.subheader("本班试卷成绩明细")
            if class_score.empty:
                st.warning("暂无学生提交该试卷")
            else:
                st.dataframe(class_score, use_container_width=True)
                # 新增：导出成绩明细Excel
                excel_file = export_score_detail_excel(class_score, selected_class, paper_map[select_test])
                st.download_button(
                    label="📥 导出成绩明细 Excel",
                    data=excel_file,
                    file_name=f"{selected_class}_{paper_map[select_test]}_成绩明细.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

        # 统计分析 + 新增Word导出按钮
        if show_stat:
            if class_score.empty:
                st.warning("暂无提交数据")
            else:
                total_num = len(class_score)
                avg = round(class_score["total_score"].mean(), 2)
                max_s = class_score["total_score"].max()
                min_s = class_score["total_score"].min()

                # 新增：计算及格率、优秀率，整理统计数据
                pass_count = len(class_score[class_score["total_score"] >= 60])
                pass_rate = round(pass_count / total_num * 100, 2)
                excellent_count = len(class_score[class_score["total_score"] >= 90])
                excellent_rate = round(excellent_count / total_num * 100, 2)
                stat_data = {
                    "total_count": total_num,
                    "avg_score": avg,
                    "max_score": max_s,
                    "min_score": min_s,
                    "pass_rate": pass_rate,
                    "excellent_rate": excellent_rate
                }

                # 新增：计算分数段分布
                bins = [0, 60, 70, 80, 90, 101]
                seg_labels = ["0-59分", "60-69分", "70-79分", "80-89分", "90-100分"]
                class_score["score_segment"] = pd.cut(class_score["total_score"], bins=bins, labels=seg_labels, right=False)
                segment_data = class_score["score_segment"].value_counts().sort_index().to_dict()

                # 新增：初始化逐题统计列表
                question_stat_list = []

                m1, m2, m3, m4 = st.columns(4)
                with m1:
                    st.metric("参与人数", total_num)
                with m2:
                    st.metric("平均分", avg)
                with m3:
                    st.metric("最高分", max_s)
                with m4:
                    st.metric("最低分", min_s)

                st.subheader("成绩等级分布")
                st.dataframe(class_score["level"].value_counts(), use_container_width=True)
                st.divider()

                st.subheader("逐题作答分析")
                for q in q_list:
                    qid = q["id"]
                    score_col = f"q{qid}"
                    ans_col = f"user_ans_{qid}"
                    st.markdown(f"### 第{qid}题（满分{q['score']}分）")
                    st.write(f"题干：{q['title']}")
                    if q["type"] in ["single", "multi", "judge"]:
                        st.write(f"选项：{' | '.join(q['opts'])}")
                    std_ans_txt = "、".join(q["ans"]) if isinstance(q["ans"], list) else q["ans"]
                    st.write(f"标准答案：{std_ans_txt}")

                    if score_col in class_score.columns:
                        q_avg_score = round(class_score[score_col].mean(), 2)
                        right_cnt = len(class_score[class_score[score_col] == q["score"]])
                        wrong_cnt = total_num - right_cnt
                        acc_rate = round(right_cnt / total_num * 100, 2)
                        st.write(f"本题平均分：{q_avg_score} | 答对{right_cnt}/{total_num} | 正确率{acc_rate}%")

                        # 新增：收集逐题统计数据
                        question_stat_list.append({
                            "题号": qid,
                            "题型": q["type"],
                            "分值": q["score"],
                            "正确率": acc_rate
                        })

                        if q["type"] in ["single", "judge"]:
                            opt_count = {"A": 0, "B": 0, "C": 0, "D": 0}
                            if ans_col in class_score.columns:
                                for a in class_score[ans_col].dropna():
                                    if a in opt_count:
                                        opt_count[a] += 1
                            bar_df = pd.DataFrame(
                                {"选项": list(opt_count.keys()), "作答人数": list(opt_count.values())})
                            st.bar_chart(bar_df, x="选项", y="作答人数", height=240)
                        else:
                            bar_df = pd.DataFrame({"结果": ["答对", "答错"], "人数": [right_cnt, wrong_cnt]})
                            st.bar_chart(bar_df, x="结果", y="人数", height=240)
                    else:
                        st.info("暂无该题作答数据")
                        # 新增：无作答数据时也记录
                        question_stat_list.append({
                            "题号": qid,
                            "题型": q["type"],
                            "分值": q["score"],
                            "正确率": 0
                        })
                    st.divider()

                # 新增：生成逐题DataFrame + 导出Word报告按钮
                question_df = pd.DataFrame(question_stat_list)
                word_file = export_paper_analysis_word(stat_data, segment_data, question_df, selected_class, paper_map[select_test])
                st.download_button(
                    label="📥 导出试卷分析报告 Word",
                    data=word_file,
                    file_name=f"{selected_class}_{paper_map[select_test]}_试卷分析报告.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

# ===================== 六、班级教学资源预览 =====================
st.divider()
st.subheader("六、班级教学资源")
if not selected_class:
    st.info("请先选择上方班级")
else:
    res_list = get_class_resources(selected_class)
    if len(res_list) == 0:
        st.info("该班级暂无分配的教学资源，请先在资源管理中上传并分配")
    else:
        # 按分类分组展示
        type_groups = {}
        for r in res_list:
            t = r["res_type"]
            if t not in type_groups:
                type_groups[t] = []
            type_groups[t].append(r)
        for res_type, items in type_groups.items():
            st.markdown(f"#### {res_type}")
            for r in items:
                with st.expander(f"📂 {r['res_name']}", expanded=False):
                    st.caption(f"上传时间：{r['create_time']}")
                    f_path = os.path.join(RESOURCE_FILE_DIR, r["file_name"])
                    ext = r["file_name"].split(".")[-1].lower()
                    show_name = r["file_name"].split("_", 1)[1] if "_" in r["file_name"] else r["file_name"]
                    if os.path.exists(f_path):
                        if ext in IMG_EXT:
                            st.image(f_path, caption=show_name, use_container_width=True)
                        elif ext in VIDEO_EXT:
                            st.video(f_path)
                        # 所有格式都提供下载按钮
                        with open(f_path, "rb") as fp:
                            file_bytes = fp.read()
                        st.download_button(
                            label=f"📥 下载文件：{show_name}",
                            data=file_bytes,
                            file_name=show_name,
                            use_container_width=True
                        )

st.divider()
st.page_link("pages/00_教师入口首页.py", label="← 教师入口首页", use_container_width=True)